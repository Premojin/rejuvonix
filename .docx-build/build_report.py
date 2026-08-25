from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\HoweT\Downloads\Dev\Rejuvonix")
OUT = ROOT / "artifacts" / "Rejuvonix_Eligibility_Flow_Discovery_and_Implementation_Plan.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "243746"
MUTED = "66727D"
LIGHT = "F2F4F7"
CALLOUT = "F4F6F9"
TEAL = "157A7C"
TEAL_LIGHT = "E7F3F3"
GOLD = "7A5A00"
GOLD_LIGHT = "FFF6D8"
RED = "9B1C1C"
RED_LIGHT = "FCEBEB"
WHITE = "FFFFFF"
TABLE_W = 9360
TABLE_IND = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table, color="C8CED5", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    assert sum(widths) == TABLE_W
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_W))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_IND))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    set_table_borders(table)


def set_font(run, size=11, color=INK, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([r_style, color_el, underline])
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.extend([r_pr, text_el])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    list_bullet = styles["List Bullet"]
    list_bullet.font.name = "Calibri"
    list_bullet._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    list_bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    list_bullet.font.size = Pt(11)
    list_bullet.font.color.rgb = RGBColor.from_string(INK)
    list_bullet.paragraph_format.space_before = Pt(0)
    list_bullet.paragraph_format.space_after = Pt(8)
    list_bullet.paragraph_format.line_spacing = 1.167


def add_numbering(doc, kind="bullet"):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    # Use a high private range so Word does not coalesce these definitions with
    # built-in template numbering during PDF export.
    abstract_id = max(abstract_ids + [99]) + 1
    num_id = max(num_ids + [99]) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"A0{abstract_id:06X}"[-8:])
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, suffix, p_pr, r_pr])
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def patch_numbering_geometry(doc, num_id):
    numbering = doc.part.numbering_part.element
    num = next(x for x in numbering.findall(qn("w:num")) if x.get(qn("w:numId")) == str(num_id))
    abstract_id = num.find(qn("w:abstractNumId")).get(qn("w:val"))
    abstract = next(x for x in numbering.findall(qn("w:abstractNum")) if x.get(qn("w:abstractNumId")) == abstract_id)
    lvl = abstract.find(qn("w:lvl"))
    p_pr = lvl.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        lvl.append(p_pr)
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        p_pr.append(tabs)
    for child in list(tabs):
        tabs.remove(child)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")


def clone_num_instance(doc, source_num_id):
    numbering = doc.part.numbering_part.element
    source = next(x for x in numbering.findall(qn("w:num")) if x.get(qn("w:numId")) == str(source_num_id))
    abstract_id = source.find(qn("w:abstractNumId")).get(qn("w:val"))
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(num_ids + [99]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def numbered_para(doc, text, num_id, bold_prefix=None):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_callout(doc, title, body, level="info"):
    fills = {"info": TEAL_LIGHT, "warning": GOLD_LIGHT, "risk": RED_LIGHT}
    colors = {"info": TEAL, "warning": GOLD, "risk": RED}
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_W])
    set_cell_shading(table.cell(0, 0), fills[level])
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_font(r, size=11, color=colors[level], bold=True)
    p2 = table.cell(0, 0).add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, font_size=9.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_font(r, size=font_size, color=DARK_BLUE, bold=True)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        row = table.add_row()
        keep_row_together(row)
        for i, value in enumerate(row_data):
            p = row.cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            set_font(r, size=font_size, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_source_line(doc, label, links):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label + ": ")
    set_font(r, size=9, color=MUTED, bold=True)
    for idx, (text, url) in enumerate(links):
        if idx:
            sep = p.add_run(" | ")
            set_font(sep, size=9, color=MUTED)
        add_hyperlink(p, text, url)


def add_page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("REJUVONIX  |  ELIGIBILITY FLOW DISCOVERY")
    set_font(r, size=8.5, color=MUTED, bold=True)
    p_border = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "D7DBE2")
    pbdr.append(bottom)
    p_border.append(pbdr)
    add_page_number(section.footer.paragraphs[0])

    bullet_id = add_numbering(doc, "bullet")
    decimal_id = add_numbering(doc, "decimal")

    # Memo masthead
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("DISCOVERY REPORT")
    set_font(r, size=10, color=TEAL, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Eligibility Assessment Flow")
    set_font(r, size=25, color=INK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Source-to-destination implementation plan and approval gate")
    set_font(r, size=14, color=MUTED)
    metadata = [
        ("Project", "Rejuvonix telehealth website"),
        ("Source", "rejuvonix.com/eligibility"),
        ("Destination", "rejuvonix.otrayoe.chatgpt.site"),
        ("Prepared", "August 25, 2026"),
        ("Status", "Plan only — implementation blocked pending approval"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        set_font(r, size=10.5, color=INK, bold=True)
        r = p.add_run(value)
        set_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    add_callout(
        doc,
        "Decision gate",
        "Do not implement the full source intake yet. It gathers health history, prescription-use details, identity/contact data, uploads, consent, and shipping information, while the destination currently has no suitable intake backend or protected-data controls. Approval is required for either a non-submitting, non-PHI pre-screen or a full intake backed by reviewed clinical, privacy, security, and operational infrastructure.",
        "risk",
    )
    doc.add_heading("Executive outcome", level=1)
    p = doc.add_paragraph(
        "The source flow can be reproduced technically, but not safely as a complete production intake using the destination’s current architecture. The destination has a three-question informational questionnaire and no intake API, protected storage, authenticated member workflow, upload service, consent ledger, or established health-data analytics policy."
    )
    numbered_para(doc, "Recommended near-term option: build a non-submitting, non-PHI eligibility pre-screen that routes qualified visitors to a clearly disclosed next step.", bullet_id)
    numbered_para(doc, "Alternative: pause implementation until the prescribing medical group, privacy counsel, security owner, and operations team approve the full data lifecycle and clinical logic.", bullet_id)
    numbered_para(doc, "Do not silently copy promotional claims, testimonials, approval language, or medical outcome logic without review.", bullet_id)

    add_page_break(doc)
    doc.add_heading("1. Scope, methods, and system context", level=1)
    doc.add_paragraph(
        "This report documents the visitor-facing source assessment, identifies assessment-related calls to action across both sites, maps the source flow to the current destination architecture, and defines the minimum decisions required before code changes. The assessment was inspected without bypassing authentication and without using real patient information."
    )
    doc.add_heading("Sites reviewed", level=2)
    add_table(doc, ["Surface", "Role", "Observed access"], [
        ("rejuvonix.com", "Approved public source website", "Public pages and eligibility flow reviewed"),
        ("rejuvonix.com/eligibility", "Approved source assessment", "Visitor-facing flow and branching documented"),
        ("rejuvonix.otrayoe.chatgpt.site", "Destination website", "External URL returned HTTP 401 during review; local project inspected"),
        ("Local Rejuvonix repository", "Destination implementation source", "Routing, components, storage, tests, and hosting configuration inspected"),
    ], [2200, 3100, 4060])
    doc.add_heading("Destination architecture", level=2)
    for item in [
        "Next.js 16 / React 19 application using the Vinext compatibility layer and Cloudflare-oriented deployment configuration.",
        "Public routes are implemented in the app directory; shared public chrome is reused by many pages, but some page-level UI remains duplicated.",
        "The only observed API surface is a health endpoint. There is no patient-intake submission API.",
        "The current /get-started experience contains three informational questions and does not persist or submit answers.",
        "The sign-in route is a visual mockup, not a completed authenticated member portal.",
        "The hosting configuration exposes no bound D1 database or R2 object storage, and the D1 schema is empty.",
        "No established analytics event system or destination-specific privacy/consent capture mechanism was found for assessment answers.",
    ]:
        numbered_para(doc, item, bullet_id)
    doc.add_heading("Public source routes crawled", level=2)
    routes = "/, /about, /contact, /disclaimer, /eligibility, /hipaa, /privacy, /terms, /waitlist, /treatments/glp-1-microdose, /treatments/semaglutide, /treatments/tirzepatide"
    doc.add_paragraph(routes)
    add_source_line(doc, "Primary sources", [
        ("Source home", "https://rejuvonix.com/"),
        ("Source eligibility", "https://rejuvonix.com/eligibility"),
        ("Destination", "https://rejuvonix.otrayoe.chatgpt.site/"),
    ])

    add_page_break(doc)
    doc.add_heading("2. Source eligibility flow map", level=1)
    doc.add_paragraph(
        "The source is a multi-stage intake with 44 configured fields, conditional questions, early lead transmission, automated disqualification rules, and a post-submission approval/checkout path. The following map separates general wellness screening from health, prescription, contact, consent, and fulfillment data."
    )
    doc.add_heading("High-level sequence", level=2)
    for step in [
        "Start eligibility → wellness and motivation questions",
        "Sex/pregnancy branch → body measurements and BMI gate",
        "Name/contact capture → early lead server call and Meta matching",
        "Medical history → surgery, GLP-1, medication, and allergy branches",
        "Risk-tolerance and promotional interstitial",
        "Date of birth and shipping address",
        "Accuracy/terms/privacy/telehealth and automated-contact consents",
        "Client disqualifier check → server submission → curating state",
        "Approved result → pharmacy promotion → treatment checkout",
    ]:
        numbered_para(doc, step, decimal_id)
    add_callout(
        doc,
        "Important sequencing risk",
        "The source makes an early lead server call, including name, email, and phone and Meta matching, as soon as those fields validate—before the visitor reaches the final consent screen. That behavior should not be copied without explicit privacy and marketing-consent approval.",
        "warning",
    )
    add_page_break(doc)
    doc.add_heading("Stage A — General eligibility", level=2)
    add_table(doc, ["Step", "Prompt / field", "Options or validation", "Classification"], [
        ("1", "Primary motivation", "Long term health; Feel better day to day; Confidence and appearance; Reduce health risks", "General screening"),
        ("2", "Current activity", "Very active; Moderately active; Not very active", "Lifestyle"),
        ("3", "Typical sleep", ">9 hours; 7–9; <7; varies/trouble", "Lifestyle / health-adjacent"),
        ("4", "Stress frequency", "Rarely; Moderately; Often", "Health-adjacent"),
        ("5", "Gender", "Male; Female", "Sensitive personal data"),
        ("6", "Pregnancy branch", "If female: No / Yes. Yes disqualifies.", "Sensitive health data"),
        ("7", "Height, current weight, goal weight", "Height 3–8 ft and 0–11 in; weights 50–600", "Health data"),
        ("8", "BMI interstitial", "BMI below 25 disqualifies; promotional outlook/timeline copy shown", "Derived health data / claim"),
        ("9", "First and last name", "Both required", "Identity data"),
        ("10", "Email and phone", "Both required and validated", "Contact data"),
    ], [620, 2240, 3900, 2600], font_size=8.6)

    add_page_break(doc)
    doc.add_heading("Stage B — Medical history and treatment context", level=2)
    add_table(doc, ["Prompt / field", "Choices / conditional behavior", "Sensitivity"], [
        ("Body-shape selection", "Four visual choices", "Body/health profile"),
        ("Ethnicity", "Six choices", "Sensitive demographic data"),
        ("Current health effects", "Multi-select: none, low libido, skin, brain fog", "Health symptoms"),
        ("Diagnoses / contraindications", "Multi-select list covering thyroid cancer history, MEN2, cancer, substance misuse, eating disorder, serious mental-health conditions, pancreatitis, diabetes, liver/kidney/gallbladder disease, and persistent hypoglycemia", "High-sensitivity medical history"),
        ("Prior surgeries", "Yes/No; details required when Yes", "Medical history"),
        ("Weight-loss surgery", "None; bypass; duodenal switch; lap band; sleeve", "Medical history"),
        ("GLP-1 history", "Never taken / Taken", "Prescription treatment history"),
        ("If GLP-1 taken", "Medication name, last injection dose, last-dose date, duration, optional side effects, photo availability/upload, data visibility confirmation, optional preferred starting dose, required dose acknowledgment", "Prescription and upload data"),
        ("Medications / supplements", "Yes/No; details required when Yes", "Medication data"),
        ("Specific allergies", "None; glycine; B12; B3; unsure", "Medical data"),
        ("Other allergies", "Yes/No; details required when Yes", "Medical data"),
        ("Resting heart rate", "Unknown; <70; 70–85; >85", "Health measurement"),
        ("Tolerance for side-effect disruption", "High; Moderate; Low", "Treatment preference"),
    ], [2300, 5000, 2060], font_size=8.5)
    doc.add_paragraph("A promotional testimonial interstitial claiming a 60% outcome appears after the risk-tolerance question.")
    doc.add_heading("Automated medical disqualifiers observed", level=2)
    for item in [
        "Medullary thyroid carcinoma (personal history)",
        "Family history of medullary thyroid carcinoma",
        "Multiple Endocrine Neoplasia type 2 (MEN2)",
        "Pancreatitis",
        "Active gallbladder disease",
    ]:
        numbered_para(doc, item, bullet_id)
    doc.add_paragraph(
        "Other serious diagnoses appear in the questionnaire but were not observed in the same client-side automated-disqualifier list. That discrepancy requires clinical review; the destination should not infer or expand medical decision logic."
    )

    add_page_break(doc)
    doc.add_heading("Stage C — Completion, consent, and outcomes", level=2)
    add_table(doc, ["Area", "Observed behavior", "Implementation concern"], [
        ("Date of birth", "Required; accepted age range 18–100", "Identity and health eligibility; verify age rules and state/service restrictions"),
        ("Shipping address", "Line 1 required; line 2 optional; city, state, and ZIP required", "Fulfillment data should not be collected without a defined operational need and protected lifecycle"),
        ("Accuracy / legal consent", "Required acknowledgment referencing Terms, Privacy, and Telehealth", "Exact approved legal text and versioned consent evidence are needed"),
        ("Automated contact consent", "Required", "Marketing/communications counsel should review channel, scope, revocation, and recordkeeping"),
        ("Submission", "Client disqualifier check, then server call, curating state, approved outcome, pharmacy promotion, treatment checkout", "Do not imply treatment or prescription is guaranteed; server-side clinical control required"),
        ("Errors", "Retry or review responses", "Preserve answers securely and avoid exposing sensitive content in logs or URLs"),
        ("Exit", "Header logo exits without confirmation", "Risk of accidental abandonment"),
        ("Refresh / return", "Answers reset; no complete persistence observed", "Clarify expected retention and restart behavior"),
        ("Back navigation", "Visitor can return and change a disqualifying answer", "Server must independently validate; client-side outcomes are not authoritative"),
    ], [1700, 3900, 3760], font_size=8.5)
    doc.add_heading("Automated outcome rules observed", level=2)
    for item in [
        "BMI below 25 → ineligible",
        "Pregnant, breastfeeding, or trying to become pregnant → ineligible",
        "Personal or family history of medullary thyroid carcinoma → ineligible",
        "MEN2 → ineligible",
        "Pancreatitis → ineligible",
        "Active gallbladder disease → ineligible",
    ]:
        numbered_para(doc, item, bullet_id)
    add_callout(
        doc,
        "Outcome-language concern",
        "The source displays “You’re approved” before checkout. A destination implementation should distinguish questionnaire screening from an independent prescriber’s medical decision and should not promise treatment or a prescription.",
        "risk",
    )

    add_page_break(doc)
    doc.add_heading("3. Source CTA inventory", level=1)
    doc.add_paragraph(
        "The source repeatedly directs visitors toward eligibility using global navigation and page-specific actions. Informational links were kept separate from patient-intake actions."
    )
    add_table(doc, ["Location", "CTA text / element", "Observed intent"], [
        ("Global site chrome", "Build My Protocol", "Open or navigate to eligibility/intake"),
        ("Global site chrome", "Build My Protocol →", "Open or navigate to eligibility/intake"),
        ("Global site chrome", "Check Eligibility", "Navigate to eligibility"),
        ("Homepage hero/content", "Get Started", "Start patient-intake journey"),
        ("Homepage cards", "Hair Restoration; Skin Regeneration; Performance; Sexual Health", "Card actions appear intake-oriented and require product/path review"),
        ("About page", "Schedule Your Consultation", "Patient-intake / consultation action"),
        ("Treatment pages", "Eligibility- or protocol-oriented actions", "Start assessment or proceed toward treatment"),
    ], [2100, 3400, 3860])
    doc.add_heading("Informational navigation", level=2)
    doc.add_paragraph(
        "Learn More, About, Contact, treatment-detail navigation, Terms, Privacy, HIPAA, disclaimer, and other policy links should remain informational unless their actual purpose clearly starts assessment. The destination should not convert general discovery links into eligibility actions merely to increase CTA consistency."
    )
    doc.add_heading("Source behavior notes", level=2)
    for item in [
        "The source mixes direct navigation, modal-like progression, promotional interstitials, medical questions, and checkout-oriented steps in one journey.",
        "CTA language varies across pages, including Build My Protocol, Get Started, Check Eligibility, and Schedule Your Consultation.",
        "Some cards use treatment-category labels that do not by themselves make the destination or next action clear.",
    ]:
        numbered_para(doc, item, bullet_id)

    add_page_break(doc)
    doc.add_heading("4. Source-to-destination screen mapping", level=1)
    add_table(doc, ["Source stage", "Destination equivalent", "Disposition"], [
        ("Eligibility start / wellness questions", "New native /eligibility route using existing typography, controls, and responsive patterns", "Suitable for an approved non-PHI pre-screen"),
        ("Sex, pregnancy, BMI", "New accessible conditional form steps", "Requires clinical review; pregnancy and BMI are health screening"),
        ("Name, email, phone", "No safe existing intake equivalent", "Do not implement early lead submission without consent and privacy approval"),
        ("Medical history / contraindications", "No existing protected medical-intake system", "Blocked for full implementation"),
        ("GLP-1 history and photo upload", "No secure upload or object-store binding", "Blocked; do not store in public client or generic logs"),
        ("DOB and shipping address", "No fulfillment backend", "Blocked unless approved operational requirement and protected lifecycle exist"),
        ("Legal / telehealth / contact consent", "Destination has policy links but no versioned consent ledger", "Needs counsel-approved language and evidence model"),
        ("Automated disqualification", "Could be reproduced in client UI, but client logic is not authoritative", "Needs server validation and clinical ownership"),
        ("Approved result and checkout", "No equivalent prescribing or checkout workflow identified", "Do not copy; define reviewed next step"),
        ("/get-started", "Existing three-question informational questionnaire", "Preserve initially; later redirect only after approved route replaces it"),
    ], [2300, 4000, 3060], font_size=8.5)
    doc.add_heading("Proposed route", level=2)
    doc.add_paragraph(
        "Use /eligibility as the canonical destination route. Preserve /get-started during implementation and testing; once the new flow is approved, either convert it to a compatibility redirect or present a deliberate handoff so existing links do not break."
    )
    doc.add_heading("Presentation approach", level=2)
    for item in [
        "Build the flow natively—do not iframe the source.",
        "Reuse destination form controls, buttons, spacing, typography, color tokens, focus treatment, and responsive containers.",
        "Use a clear step title, visible progress, back/continue controls, inline validation, and a restart/exit pattern with confirmation when answers would be lost.",
        "Preserve answers when moving backward, subject to the approved privacy model.",
        "Keep mobile controls reachable without horizontal overflow and announce step and validation changes to assistive technology.",
    ]:
        numbered_para(doc, item, bullet_id)

    add_page_break(doc)
    doc.add_heading("5. Destination CTAs", level=1)
    doc.add_heading("Assessment-related CTAs to connect after approval", level=2)
    add_table(doc, ["Destination area", "Likely component / file", "Connection rule"], [
        ("Homepage eligibility actions and modal result", "app/page.tsx", "Route intentional assessment actions to /eligibility; preserve modal accessibility work and eligibility logic unless separately approved"),
        ("Shared header and footer", "Shared SiteChrome component", "Connect only Get Started / Check Eligibility / assessment-oriented labels; leave informational navigation unchanged"),
        ("Shared product/treatment CTA", "ProductPage component", "Connect assessment actions while preserving treatment detail navigation"),
        ("Wegovy experience CTA", "WegovyInjectionExperience component", "Connect only the assessment-oriented action"),
        ("Compounded, how-it-works, safety, support, treatments, sign-in", "Corresponding app routes/components", "Audit each CTA by visible intent; connect assessment actions, retain informational and authentication actions"),
    ], [2500, 2800, 4060], font_size=8.8)
    doc.add_heading("CTAs that should remain unchanged", level=2)
    for item in [
        "Learn More and treatment-detail links that lead to explanatory content.",
        "Header/footer About, Contact, treatment navigation, policy, privacy, HIPAA, terms, and disclaimer links.",
        "Sign-in controls, which should not be redirected to eligibility.",
        "Any support or contact link whose visible purpose is assistance rather than assessment.",
        "Carousel, BMI calculator, treatment cards, and other unrelated interactions unless a specific control is explicitly assessment-oriented.",
    ]:
        numbered_para(doc, item, bullet_id)
    add_callout(
        doc,
        "Consistency rule",
        "Use one approved assessment label site-wide where practical, but do not change informational labels or clinical/product wording solely for uniformity. Every connected control should land on the first assessment screen, not a mid-flow state.",
        "info",
    )

    add_page_break(doc)
    doc.add_heading("6. Proposed data-handling approach", level=1)
    doc.add_heading("Option A — Non-submitting pre-screen (recommended first)", level=2)
    for item in [
        "Limit questions to the minimum approved eligibility indicators and avoid name, contact, free-text medical history, prescription details, uploads, DOB, and shipping address.",
        "Keep answers in component memory only; do not place answers in URLs, analytics, console output, error reports, or persistent browser storage.",
        "Provide clear screening language: the result is informational and does not guarantee treatment, prescribing, or fulfillment.",
        "Route visitors to a reviewed next step without transmitting health answers.",
        "Use synthetic data only in tests and screenshots.",
    ]:
        numbered_para(doc, item, bullet_id)
    doc.add_heading("Option B — Full source-equivalent intake", level=2)
    doc.add_paragraph("Before implementation, the project needs all of the following:")
    requirements = [
        ("Clinical governance", "Named clinical owner; approved question set; disqualifier and outcome logic; prescriber-independent decision path; emergency/urgent-care limitations"),
        ("Backend", "Authenticated or otherwise protected intake API; strict server validation; idempotent submission; controlled error handling"),
        ("Storage", "Encryption in transit and at rest; defined database records; secure file uploads; least-privilege service bindings"),
        ("Access control", "Role-based access, staff authentication, audit logging, and separation of operational and clinical permissions"),
        ("Privacy lifecycle", "Data minimization; notice; purpose limitation; retention/deletion schedule; access/correction/deletion workflow; vendor inventory"),
        ("Consent", "Counsel-approved telehealth, privacy, terms, and automated-contact language with version, timestamp, source, and revocation handling"),
        ("Analytics", "Explicit policy prohibiting health answers, contact fields, uploads, and derived eligibility outcomes from general analytics/advertising payloads"),
        ("Operations", "Eligible states/services, prescribing medical group, pharmacy/fulfillment handoff, support and escalation procedures"),
        ("Security", "Threat model, log review, secrets management, abuse controls, incident response, and vendor agreements where required"),
    ]
    add_table(doc, ["Capability", "Minimum requirement"], requirements, [2200, 7160], font_size=8.8)
    doc.add_heading("Prohibited client-side handling", level=2)
    for item in [
        "No sensitive answers in query strings, hashes, page titles, referrers, client logs, or analytics event properties.",
        "No persistent localStorage/sessionStorage cache for a full medical intake without explicit review.",
        "No direct upload to a public or unprotected bucket.",
        "No client-only determination that is represented as clinical approval.",
        "No early marketing lead transmission before a valid, reviewed consent basis.",
    ]:
        numbered_para(doc, item, bullet_id)

    add_page_break(doc)
    doc.add_heading("7. Expected implementation surface", level=1)
    doc.add_paragraph(
        "Exact files should be confirmed immediately before implementation because the repository is active. The expected surface below is deliberately narrow and does not authorize changes."
    )
    add_table(doc, ["Area", "Expected change", "Notes"], [
        ("New eligibility route", "app/eligibility/page.tsx and narrowly scoped local components/tests", "Native flow and accessible step management"),
        ("Shared assessment components", "Reuse or add within existing component conventions", "Avoid a second design system or duplicated chrome"),
        ("Homepage", "app/page.tsx", "Connect approved assessment CTAs only; avoid unrelated modal/carousel changes"),
        ("Shared navigation", "SiteChrome component", "Connect approved header/footer assessment controls"),
        ("Product pages", "ProductPage and specific experience component", "Only assessment-oriented actions"),
        ("Current questionnaire", "app/get-started/page.tsx", "Compatibility route or handoff after approval"),
        ("Tests", "Existing contract/UI tests or narrowly scoped additions", "Cover route, branching, keyboard, validation, and no sensitive telemetry"),
        ("Backend/config", "Not defined", "Stop if full intake is approved until backend/security design is separately authorized"),
    ], [2100, 3400, 3860], font_size=8.8)
    doc.add_heading("Areas intentionally outside scope", level=2)
    for item in [
        "Unrelated destination page redesigns, typography, navigation structure, treatment content, pricing, product claims, authentication, database, worker, or deployment behavior.",
        "Source website changes or private/patient systems.",
        "Generated replacements for real product photography.",
        "Claims editing without an authoritative source and appropriate owner review.",
        "Production deployment, domain changes, merge, or publication.",
    ]:
        numbered_para(doc, item, bullet_id)

    add_page_break(doc)
    doc.add_heading("8. Compliance and visitor-understanding review", level=1)
    doc.add_paragraph(
        "This is a technical and content-risk review, not legal advice or a compliance certification. Confirmed observations are separated from decisions requiring healthcare counsel, privacy counsel, the prescribing medical group, pharmacy, or marketing review."
    )
    doc.add_heading("Claims and representations requiring review", level=2)
    add_table(doc, ["Observed source representation", "Why it needs review", "Owner"], [
        ("“Lose weight without cravings—starts this week”", "Outcome and timing claim; may imply predictable individual response", "Clinical + marketing/legal"),
        ("“92% success outlook”", "Quantified success representation; substantiation and definition of success needed", "Clinical + marketing/legal"),
        ("“−18% body weight” and goal timeline", "Quantified efficacy and projected personal timeline", "Clinical + marketing/legal"),
        ("“Down 38 lbs,” real-member testimonials, before/after material", "Endorsement, typicality, disclosure, and image authenticity concerns", "Marketing/legal"),
        ("“You’re approved”", "May blur screening, medical eligibility, and independent prescribing decision", "Clinical + legal"),
        ("“Safest, highest-quality,” testing every lot, “All dosages”", "Comparative/superlative, quality, testing, and availability claims need precise substantiation", "Pharmacy + quality + legal"),
        ("Broad HIPAA representations", "Applicability and actual safeguards depend on entities, data flows, vendors, and contracts", "Privacy/security counsel"),
    ], [3000, 4200, 2160], font_size=8.5)
    doc.add_heading("Prescription and compounded-drug boundaries", level=2)
    for item in [
        "Clearly distinguish FDA-approved branded products from compounded preparations whenever either is discussed.",
        "Do not state or imply that compounded drugs are FDA-approved.",
        "Do not promise a prescription or imply the prescriber’s independent decision is predetermined.",
        "Where dosing language appears, use: “Your prescriber will determine your dose and treatment schedule based on your individual medical needs.”",
        "Do not add efficacy, safety, weight-loss, anti-aging, hormone, sexual-health, hair-restoration, or skin-restoration claims without current authoritative substantiation and documentation.",
    ]:
        numbered_para(doc, item, bullet_id)
    add_source_line(doc, "Authoritative review sources", [
        ("FDA: Compounding and FDA Q&A", "https://www.fda.gov/drugs/human-drug-compounding/compounding-and-fda-questions-and-answers"),
        ("FDA: Concerns with unapproved GLP-1 drugs", "https://www.fda.gov/drugs/postmarket-drug-safety-information-patients-and-providers/fdas-concerns-unapproved-glp-1-drugs-used-weight-loss"),
        ("FTC: Health Products Compliance Guidance", "https://www.ftc.gov/business-guidance/resources/health-products-compliance-guidance"),
        ("FTC: Endorsement Guides Q&A", "https://www.ftc.gov/business-guidance/resources/ftcs-endorsement-guides-what-people-are-asking"),
    ])

    add_page_break(doc)
    doc.add_heading("9. Accessibility and interaction requirements", level=1)
    doc.add_paragraph("Any approved implementation should meet WCAG 2.2 Level AA where applicable and preserve the destination’s established visual system.")
    add_table(doc, ["Interaction", "Required behavior"], [
        ("Step changes", "Move focus to the new question heading/group and announce step/progress without duplicate ARIA output"),
        ("Keyboard", "Logical tab order; all choices and controls operable; no keyboard traps; visible focus"),
        ("Validation", "Programmatic labels, field instructions, clear inline errors, error summary/focus where appropriate"),
        ("Back/restart/exit", "Preserve answers on back; confirm destructive restart/exit; return focus predictably"),
        ("Conditional fields", "Reveal programmatically, announce when needed, and remove hidden required states"),
        ("Mobile", "No horizontal overflow; reachable controls and safe-area spacing at 390px and 320px"),
        ("Motion", "Honor prefers-reduced-motion and avoid unnecessary transitions"),
        ("Sensitive content", "No answer content in accessible names, URLs, logs, or telemetry beyond what the user must hear on-screen"),
        ("Outcome", "Explain that screening does not guarantee treatment; provide a safe next step and a way to review answers"),
    ], [2200, 7160], font_size=8.9)
    doc.add_heading("Browser-test status", level=2)
    add_callout(
        doc,
        "Visual confirmation remains pending",
        "The external destination returned HTTP 401 during review, and prior browser automation did not have an active browser. Source-level findings are complete, but desktop/mobile screenshots and visual behavior at 1440, 1024, 768, 390, and 320 pixels must be captured during implementation verification.",
        "warning",
    )

    add_page_break(doc)
    doc.add_heading("10. Decisions and approval required", level=1)
    doc.add_paragraph("Choose one implementation boundary before any eligibility code is added:")
    add_table(doc, ["Decision", "Scope", "Readiness", "Recommendation"], [
        ("A. Non-submitting pre-screen", "Minimum approved eligibility indicators; no identity/contact, free-text medical history, uploads, DOB, or address", "Can proceed after clinical/copy approval and CTA confirmation", "Recommended first"),
        ("B. Full source-equivalent intake", "Medical history, prescription details, identity/contact, uploads, consent, address, outcome and handoff", "Blocked until backend, security, consent, clinical, state/service, pharmacy, and privacy requirements are approved", "Do not begin yet"),
    ], [2000, 3300, 2860, 1200], font_size=8.6)
    doc.add_heading("Approval questions", level=2)
    approval_decimal_id = clone_num_instance(doc, decimal_id)
    questions = [
        "Should the first implementation be Option A, the non-submitting non-PHI pre-screen, or should work pause for Option B infrastructure and governance?",
        "Which destination CTA labels and page locations are explicitly approved to start assessment?",
        "What is the reviewed next step for visitors who pass or fail the pre-screen?",
        "Who owns approval of the medical question set and disqualifier logic?",
        "Which states and services are in scope, and how should unavailable services be communicated?",
        "What privacy, telehealth, terms, and automated-contact language is approved, and how must consent evidence be retained?",
        "Which source promotional interstitials, claims, testimonials, before/after material, pricing, and pharmacy representations—if any—are approved to copy?",
    ]
    for q in questions:
        numbered_para(doc, q, approval_decimal_id)
    doc.add_heading("Recommended next step", level=2)
    add_callout(
        doc,
        "Approve a narrow implementation plan",
        "Authorize a native /eligibility pre-screen that keeps answers in memory, collects no identity or contact information, sends no health answers to analytics or a backend, uses clearly non-clinical outcome language, and connects only the explicitly approved assessment CTAs. Continue to hold the full intake until its protected-data and clinical workflow are designed and reviewed.",
        "info",
    )
    doc.add_heading("Record of intentional exclusions", level=2)
    for item in [
        "No code was changed as part of the discovery pass.",
        "No real patient data was used or submitted.",
        "No private systems, authenticated patient areas, credentials, environment secrets, or backend data were accessed.",
        "No production site, domain, deployment, branch, pull request, or source website was modified.",
        "No compliance certification is made by this report.",
    ]:
        numbered_para(doc, item, bullet_id)

    # Core properties; privacy scrub will clear author fields afterward.
    doc.core_properties.title = "Rejuvonix Eligibility Flow Discovery and Implementation Plan"
    doc.core_properties.subject = "Source-flow mapping, destination CTA inventory, data-handling risks, and approval gate"
    doc.core_properties.keywords = "Rejuvonix, eligibility, assessment, accessibility, privacy, implementation plan"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = "Prepared from the completed read-only discovery report."
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
